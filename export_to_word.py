from docx import Document
from docx.shared import Pt

def export_rounds_to_word(text_output, division, filename="QuizRounds.docx"):
    doc = Document()

    lines = text_output.split("\n")
    first_round_seen = False

    for line in lines:
        stripped = line.strip()

        if stripped == "":
            doc.add_paragraph()
            continue

        is_round_title = stripped.startswith(f"{division} Round")
        is_elim_title = stripped.startswith(f"{division} ") and any(
            t in stripped for t in ["Finals", "Semifinals", "Quarterfinals"]
        )
        is_title_page = stripped == f"{division} Division Questions"

        if is_title_page:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(24)
            doc.add_page_break()
            continue

        if is_round_title or is_elim_title:
            if first_round_seen:
                doc.add_page_break()
            first_round_seen = True
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(18)
            continue

        if stripped == "HALFTIME":
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            continue

        if stripped == "===PAGEBREAK===":
            doc.add_page_break()
            continue

        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(12)

    doc.save(filename)
